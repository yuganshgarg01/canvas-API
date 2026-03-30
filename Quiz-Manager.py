import requests
import time
import os
import re
import subprocess
import sys

# ============================================================
#  CONFIGURATION — Fill these in
# ============================================================
CANVAS_URL = "https://canvas.instructure.com"
API_TOKEN  = "paste_your_token_here"
COURSE_ID  = "your course ID here"
# ============================================================

headers = {"Authorization": f"Bearer {API_TOKEN}"}


# ============================================================
# STEP 0: DOWNLOAD SAMPLE TEMPLATE
# ============================================================
def download_sample_template():
    """Generate and save the sample question DOCX template."""
    print("\n📥 Generating sample template...")

    # Check if node and generate script exist
    if not os.path.exists("generate_quiz_template.js"):
        print("❌ generate_quiz_template.js not found in current folder!")
        print("👉 Please place generate_quiz_template.js in the same folder.")
        return

    result = subprocess.run(["node", "generate_quiz_template.js"], capture_output=True, text=True)
    if result.returncode == 0:
        print("✅ Sample template saved as: Quiz_Question_Template.docx")
        print("👉 Open it, fill in your questions, save and run option 2!")
    else:
        print(f"❌ Error generating template: {result.stderr}")


# ============================================================
# STEP 1: READ DOCX FILE
# ============================================================
def read_docx(filepath):
    """Extract text from docx using python-docx."""
    try:
        from docx import Document
    except ImportError:
        print("📦 Installing python-docx...")
        subprocess.run([sys.executable, "-m", "pip", "install", "python-docx", "--break-system-packages", "-q"])
        from docx import Document

    doc = Document(filepath)
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # Also read from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in lines:
                    lines.append(text)

    return lines


# ============================================================
# STEP 2: PARSE QUIZ DETAILS & QUESTIONS FROM LINES
# ============================================================
def parse_quiz_file(lines):
    """Parse quiz details and questions from extracted text lines."""

    quiz_details = {
        "title":        "Canvas Quiz",
        "description":  "",
        "time_limit":   30,
        "attempts":     1,
        "available_from": None,
        "available_until": None,
        "shuffle":      True,
    }

    questions = []

    # Parse quiz details
    for line in lines:
        l = line.lower()
        if "quiz title" in l:
            quiz_details["title"] = line.split(":", 1)[-1].strip()
        elif "quiz description" in l:
            quiz_details["description"] = line.split(":", 1)[-1].strip()
        elif "time limit" in l:
            val = re.search(r'\d+', line)
            if val:
                quiz_details["time_limit"] = int(val.group())
        elif "number of attempts" in l:
            val = re.search(r'\d+', line)
            if val:
                quiz_details["attempts"] = int(val.group())
        elif "available from" in l:
            val = re.search(r'\d{4}-\d{2}-\d{2}', line)
            if val:
                quiz_details["available_from"] = val.group() + "T00:00:00"
        elif "available until" in l:
            val = re.search(r'\d{4}-\d{2}-\d{2}', line)
            if val:
                quiz_details["available_until"] = val.group() + "T23:59:00"
        elif "shuffle" in l:
            quiz_details["shuffle"] = "yes" in l.lower()

    # Parse questions
    current_q = None
    for line in lines:
        # Detect question line: Q1. or 1. format
        q_match = re.match(r'^Q?(\d+)[.)]\s+(.+)', line, re.IGNORECASE)
        if q_match and len(line) > 10:
            if current_q:
                questions.append(current_q)
            current_q = {
                "number":  int(q_match.group(1)),
                "text":    q_match.group(2).strip(),
                "options": {},
                "answer":  "",
                "marks":   1,
                "topic":   "",
            }
            continue

        if current_q:
            # Options: a) or a.
            opt_match = re.match(r'^([abcd])[).]\s+(.+)', line, re.IGNORECASE)
            if opt_match:
                current_q["options"][opt_match.group(1).lower()] = opt_match.group(2).strip()
                continue

            # Answer
            if line.lower().startswith("answer"):
                ans = re.search(r'[abcd]', line, re.IGNORECASE)
                if ans:
                    current_q["answer"] = ans.group().lower()
                continue

            # Marks
            if line.lower().startswith("marks"):
                val = re.search(r'\d+', line)
                if val:
                    current_q["marks"] = int(val.group())
                continue

            # Topic
            if line.lower().startswith("topic"):
                current_q["topic"] = line.split(":", 1)[-1].strip()
                continue

    if current_q:
        questions.append(current_q)

    return quiz_details, questions


# ============================================================
# STEP 3: CREATE QUIZ IN CANVAS
# ============================================================
def create_quiz(quiz_details):
    """Create a quiz in Canvas and return its ID."""
    payload = {
        "quiz": {
            "title":                  quiz_details["title"],
            "description":            quiz_details["description"],
            "quiz_type":              "assignment",
            "time_limit":             quiz_details["time_limit"],
            "allowed_attempts":       quiz_details["attempts"],
            "shuffle_answers":        True,
            "shuffle_questions":      quiz_details["shuffle"],
            "published":              False,  # Keep unpublished until all questions added
            "show_correct_answers":   True,
        }
    }

    if quiz_details.get("available_from"):
        payload["quiz"]["unlock_at"] = quiz_details["available_from"]
    if quiz_details.get("available_until"):
        payload["quiz"]["lock_at"] = quiz_details["available_until"]
        payload["quiz"]["due_at"]  = quiz_details["available_until"]

    response = requests.post(
        f"{CANVAS_URL}/api/v1/courses/{COURSE_ID}/quizzes",
        headers=headers,
        json=payload
    )

    if response.status_code in [200, 201]:
        quiz = response.json()
        print(f"✅ Quiz created: {quiz['title']} (ID: {quiz['id']})")
        return quiz["id"]
    else:
        print(f"❌ Failed to create quiz: {response.status_code} → {response.text}")
        return None


# ============================================================
# STEP 4: ADD QUESTIONS TO QUIZ
# ============================================================
def add_question(quiz_id, question):
    """Add a single MCQ question to the quiz."""
    answers = []
    for key in ["a", "b", "c", "d"]:
        text = question["options"].get(key, f"Option {key.upper()}")
        answers.append({
            "answer_text":   text,
            "answer_weight": 100 if key == question["answer"] else 0,
            "answer_comments": "Correct!" if key == question["answer"] else "",
        })

    payload = {
        "question": {
            "question_name":  f"Question {question['number']}",
            "question_text":  question["text"],
            "question_type":  "multiple_choice_question",
            "points_possible": question["marks"],
            "answers":        answers,
        }
    }

    response = requests.post(
        f"{CANVAS_URL}/api/v1/courses/{COURSE_ID}/quizzes/{quiz_id}/questions",
        headers=headers,
        json=payload
    )

    if response.status_code in [200, 201]:
        print(f"  ✅ Q{question['number']}: {question['text'][:55]}...")
        return True
    else:
        print(f"  ❌ Q{question['number']} Failed: {response.text}")
        return False


# ============================================================
# STEP 5: PUBLISH QUIZ
# ============================================================
def publish_quiz(quiz_id):
    """Publish the quiz after all questions are added."""
    response = requests.put(
        f"{CANVAS_URL}/api/v1/courses/{COURSE_ID}/quizzes/{quiz_id}",
        headers=headers,
        json={"quiz": {"published": True}}
    )
    if response.status_code in [200, 201]:
        print(f"\n🚀 Quiz published successfully!")
    else:
        print(f"\n⚠️  Could not auto-publish: {response.text}")
        print("👉 Manually publish it from Canvas.")


# ============================================================
# MAIN FLOW
# ============================================================
def upload_quiz():
    """Main function to read docx and upload quiz to Canvas."""

    print("\n" + "="*55)
    print("   📝 CANVAS QUIZ UPLOADER")
    print("="*55)

    # Get file path
    filepath = input("\n📂 Enter path to your question DOCX file\n   (or press Enter for 'Quiz_Question_Template.docx'): ").strip()
    if not filepath:
        filepath = "Quiz_Question_Template.docx"

    if not os.path.exists(filepath):
        print(f"❌ File not found: {filepath}")
        return

    # Read and parse
    print(f"\n📖 Reading: {filepath}...")
    lines = read_docx(filepath)
    quiz_details, questions = parse_quiz_file(lines)

    if not questions:
        print("❌ No questions found! Check your document format.")
        return

    # Show summary
    print("\n" + "="*55)
    print("📋 QUIZ SUMMARY — Review Before Uploading")
    print("="*55)
    print(f"  Title          : {quiz_details['title']}")
    print(f"  Description    : {quiz_details['description']}")
    print(f"  Time Limit     : {quiz_details['time_limit']} minutes")
    print(f"  Attempts       : {quiz_details['attempts']}")
    print(f"  Available From : {quiz_details.get('available_from', 'Not set')}")
    print(f"  Available Until: {quiz_details.get('available_until', 'Not set')}")
    print(f"  Shuffle        : {quiz_details['shuffle']}")
    print(f"  Total Questions: {len(questions)}")
    total_marks = sum(q["marks"] for q in questions)
    print(f"  Total Marks    : {total_marks}")

    print(f"\n  Questions Preview:")
    for q in questions[:3]:
        print(f"  Q{q['number']}. {q['text'][:60]}... [{q['marks']} mark(s)]")
    if len(questions) > 3:
        print(f"  ... and {len(questions) - 3} more questions")

    # Confirm
    print("\n" + "="*55)
    confirm = input("Type 'YES' to upload this quiz to Canvas: ").strip()
    if confirm != "YES":
        print("❌ Cancelled.")
        return

    # Create quiz
    print("\n" + "="*55)
    print("🚀 Uploading to Canvas...")
    print("="*55 + "\n")

    quiz_id = create_quiz(quiz_details)
    if not quiz_id:
        return

    # Add questions
    print(f"\n📝 Adding {len(questions)} questions...\n")
    success = 0
    failed  = 0

    for q in questions:
        result = add_question(quiz_id, q)
        if result:
            success += 1
        else:
            failed += 1
        time.sleep(0.4)

    # Publish
    publish_choice = input("\n📢 Publish quiz now? (YES/NO): ").strip()
    if publish_choice.upper() == "YES":
        publish_quiz(quiz_id)
    else:
        print("👉 Quiz saved as draft. Publish it manually from Canvas when ready.")

    # Final summary
    print("\n" + "="*55)
    print(f"🎉 DONE!")
    print(f"   ✅ Questions uploaded : {success}")
    print(f"   ❌ Failed             : {failed}")
    print(f"   🔗 Quiz ID            : {quiz_id}")
    print(f"   🌐 View at: {CANVAS_URL}/courses/{COURSE_ID}/quizzes/{quiz_id}")
    print("="*55)


# ============================================================
# MENU
# ============================================================
print("\n" + "="*55)
print("   📝 CANVAS QUIZ MANAGER")
print("="*55)
print(f"   Course ID : {COURSE_ID}")
print(f"   Canvas URL: {CANVAS_URL}")
print("="*55)
print("1. Download Sample Question Template (DOCX)")
print("2. Upload Quiz from DOCX file to Canvas")
print("="*55)

choice = input("Enter choice (1/2): ").strip()

if choice == "1":
    download_sample_template()
elif choice == "2":
    upload_quiz()
else:
    print("❌ Invalid choice!")
